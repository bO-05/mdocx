import os
import markdown2
from docx import Document
from bs4 import BeautifulSoup

def markdown_to_html(md_content):
    """Convert markdown text to HTML."""
    html_content = markdown2.markdown(md_content)
    return html_content

def html_to_docx(html_content, docx_filename):
    """Convert HTML content to a DOCX file."""
    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()

    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li']):
        if element.name.startswith('h'):
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level=level)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name in ['ul', 'ol']:
            list_type = 'bullet' if element.name == 'ul' else 'number'
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style=f'List {list_type.capitalize()}')

    doc.save(docx_filename)

def process_markdown_files(folder_path):
    """Process all markdown files in the folder and convert them to DOCX."""
    for filename in os.listdir(folder_path):
        if filename.endswith('.md'):
            md_filepath = os.path.join(folder_path, filename)
            docx_filepath = os.path.join(folder_path, filename.replace('.md', '.docx'))

            if not os.path.exists(docx_filepath):
                with open(md_filepath, 'r', encoding='utf-8') as md_file:
                    md_content = md_file.read()
                    html_content = markdown_to_html(md_content)
                    html_to_docx(html_content, docx_filepath)
                    print(f'Converted {md_filepath} to {docx_filepath}')
            else:
                print(f'DOCX file for {md_filepath} already exists')

if __name__ == '__main__':
    folder_path = 'files'  # Folder containing the markdown files
    process_markdown_files(folder_path)
